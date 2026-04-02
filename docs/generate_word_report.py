#!/usr/bin/env python3
"""
Generate DG_Analysis_Report.docx — Word version of the DG Sample Code Check analysis report.

Uses python-docx for document creation and mermaid.ink for diagram rendering.
"""

import base64
import io
import os
import urllib.request
import urllib.error
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def style_table(table, header_color="1E40AF", alt_color="F1F5F9"):
    """Apply professional styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, header_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
            elif i % 2 == 0:
                set_cell_shading(cell, alt_color)


def add_mermaid_diagram(doc, mermaid_code, caption="", width=Inches(6)):
    """Fetch a Mermaid diagram as PNG from mermaid.ink and insert into the doc."""
    encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("utf-8")
    url = f"https://mermaid.ink/img/{encoded}"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        resp = urllib.request.urlopen(req, timeout=30)
        image_data = resp.read()
        stream = io.BytesIO(image_data)
        doc.add_picture(stream, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True
    except Exception as e:
        p = doc.add_paragraph(f"[Diagram could not be rendered: {e}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p2 = doc.add_paragraph(caption)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_block(doc, code_text):
    """Add a code block with monospace font and shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    # Apply shading to the paragraph
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'E2E8F0',
        qn('w:val'): 'clear',
    })
    pPr.append(shd)


# ---------------------------------------------------------------------------
# Main Report Builder
# ---------------------------------------------------------------------------

def build_report():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # ======================================================================
    # TITLE PAGE
    # ======================================================================
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DG Sample Code Check")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("System Analysis Report")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    datep = doc.add_paragraph()
    datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = datep.add_run(f"Generated {datetime.now().strftime('%B %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Healthcare Benefits Correspondence Automation\nPick/MultiValue (UniData) Platform")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # ======================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ======================================================================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report provides a comprehensive analysis of the DG Sample Code Check repository, "
        "a healthcare benefits correspondence automation system built on the Pick/MultiValue (UniData) "
        "platform. The system manages three categories of regulatory notices:"
    )

    bullets = [
        "COBRA Notices — Initial COBRA election notices for employees and dependents losing group health coverage",
        "HIPAA Privacy Notices — Notice of Privacy Practices required under HIPAA regulations",
        "Disabled Dependent Letters — Six letter types for the disabled dependent validation workflow "
        "(review requests, approvals, denials, recertifications, no-response)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph("The system comprises 3 main programs:")

    t = doc.add_table(rows=4, cols=3)
    t.rows[0].cells[0].text = "Program"
    t.rows[0].cells[1].text = "Purpose"
    t.rows[0].cells[2].text = "Letter Count"
    data = [
        ("COBRAN", "COBRA initial notice generation via SmartComm", "2"),
        ("HIPAA.NPP", "HIPAA Notice of Privacy Practices (SF via SmartComm, FI via PCL)", "1"),
        ("PRINT.DISABLED.DEP.LETTERS", "Disabled dependent letter generation via SmartComm", "6"),
    ]
    for i, (a, b, c) in enumerate(data, 1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
        t.rows[i].cells[2].text = c
    style_table(t)

    doc.add_paragraph(
        "These programs interact with 10+ shared data record types including EMPLOYEES, EMPLOYERS, "
        "EROPTIONS, DEPENDENTS, and ERADJ. All three programs submit correspondence through the "
        "*CREATE.CORRESP.TRACKING subroutine using the BASIC.SMARTCOMM.DOC.1 correspondence input definition."
    )
    doc.add_paragraph(
        "Key findings include 8 coupling points where functional concerns are tangled across the programs, "
        "and a proposed decoupling strategy with 8 extracted shared subroutines that would reduce code "
        "duplication and improve maintainability."
    )

    # ======================================================================
    # SECTION 2: RECORD TYPES
    # ======================================================================
    doc.add_heading("2. Record Types — Full Analysis", level=1)

    # 2.1 EMPLOYEES
    doc.add_heading("2.1 EMPLOYEES Record", level=2)
    doc.add_paragraph("Description: Core employee/member record containing enrollment status, coverage, effective dates, and demographic data.")
    doc.add_paragraph("Key: EEID (Employee ID)")
    doc.add_paragraph("Loaded via: $INCLUDE BP EMPLOYEES.INS → MAT EEREC")
    doc.add_paragraph("Used by: COBRAN, HIPAA.NPP, PRINT.DISABLED.DEP.LETTERS")

    t = doc.add_table(rows=12, cols=2)
    t.rows[0].cells[0].text = "Field Name"
    t.rows[0].cells[1].text = "Purpose"
    ee_fields = [
        ("EESTATUS", "Multi-valued status codes: A/P/T/D/W/R/I"),
        ("EENEED", "Trigger flag — C (COBRA), N (HIPAA NPP), X (Exclude)"),
        ("EEFILENO", "Foreign key to EMPLOYERS record (multi-valued)"),
        ("EEEFFDATE", "Multi-valued effective dates, sorted descending"),
        ("EERELATIONSHIP", "E (Employee), S (Spouse), C (Child)"),
        ("EETERMDATE", "Termination date"),
        ("EEMEDCOV/EEDENCOV/EEVISCOV", "Medical, dental, vision coverage codes"),
        ("EEHIPAA.SENT", "Dates when HIPAA notices were sent"),
        ("EEDEPNO", "Dependent number (multi-valued)"),
        ("EELAST/EEFIRST", "Employee last/first name"),
        ("EEADD1-EEZIP", "Employee mailing address"),
    ]
    for i, (f, p) in enumerate(ee_fields, 1):
        t.rows[i].cells[0].text = f
        t.rows[i].cells[1].text = p
    style_table(t)

    add_code_block(doc,
        "CALL *READREC('',MAT EEREC,EMPLOYEES,EEID,'')\n"
        "LOCATE TODAY IN EEEFFDATE<1,1> BY 'DR' SETTING EEIDX ELSE NULL\n"
        "IF EEEFFDATE<1,EEIDX> = '' AND EEIDX > 1 THEN EEIDX -= 1"
    )

    # 2.2 EMPLOYERS
    doc.add_heading("2.2 EMPLOYERS Record", level=2)
    doc.add_paragraph("Description: Employer/group record with billing configuration, state jurisdiction, trust type, and termination data. Cached in memory when FILENO changes.")
    doc.add_paragraph("Key: FILENO (from EEFILENO)")
    doc.add_paragraph("Loaded via: $INCLUDE BP EMPLOYERS.INS → MAT ERREC")
    doc.add_paragraph("Used by: COBRAN, HIPAA.NPP, PRINT.DISABLED.DEP.LETTERS")

    t = doc.add_table(rows=8, cols=2)
    t.rows[0].cells[0].text = "Field Name"
    t.rows[0].cells[1].text = "Purpose"
    er_fields = [
        ("EREFFDATE", "Multi-valued effective dates for employer configuration"),
        ("ERCOBRAFEE", "COBRA fee flag/amount — gates whether COBRA notices are sent"),
        ("ERBILLEDTO", "Last billed-through date for ERADJ billing period calculation"),
        ("ERACTUAL.STATE", "State jurisdiction for compliance rules"),
        ("ERTRUST", "Trust type — determines SF vs FI routing in HIPAA.NPP"),
        ("ERTERMDATE", "Employer termination date"),
        ("ERINS.CARRIER", "Insurance carrier code — used for FI HIPAA form lookup"),
    ]
    for i, (f, p) in enumerate(er_fields, 1):
        t.rows[i].cells[0].text = f
        t.rows[i].cells[1].text = p
    style_table(t)

    add_code_block(doc,
        "IF EEFILENO<1,EEIDX> <> FILENO THEN\n"
        "   FILENO = EEFILENO<1,EEIDX>\n"
        "   CALL *READREC('',MAT ERREC,EMPLOYERS,FILENO,'')\n"
        "END"
    )

    # 2.3 EROPTIONS
    doc.add_heading("2.3 EROPTIONS Record", level=2)
    doc.add_paragraph("Description: Employer options record containing fee schedules and feature flags. Always read alongside EMPLOYERS.")
    doc.add_paragraph("Key: FILENO (same as EMPLOYERS)")
    doc.add_paragraph("Loaded via: $INCLUDE BP EROPTIONS.INS → MAT EOREC")
    doc.add_paragraph("Used by: COBRAN, HIPAA.NPP")

    t = doc.add_table(rows=5, cols=2)
    t.rows[0].cells[0].text = "Field Name"
    t.rows[0].cells[1].text = "Purpose"
    eo_fields = [
        ("EOEFFDATE", "Multi-valued effective dates for options"),
        ("EOCOBRA.NOTICE", "Dollar amount charged per COBRA notice"),
        ("EOHIPAAP", "HIPAA privacy fee/flag — blank means skip SF billing"),
        ("EODDV", "Disabled dependent validation flag (Y/N)"),
    ]
    for i, (f, p) in enumerate(eo_fields, 1):
        t.rows[i].cells[0].text = f
        t.rows[i].cells[1].text = p
    style_table(t)

    # 2.4 DEPENDENTS
    doc.add_heading("2.4 DEPENDENTS Record", level=2)
    doc.add_paragraph("Description: Dependent-level record tracking disabled dependent letter workflow state.")
    doc.add_paragraph("Key: EEID.DEPNO (e.g., 12345.01)")
    doc.add_paragraph("Loaded via: $INCLUDE BP DEPENDENTS.INS → MAT DPREC")
    doc.add_paragraph("Used by: HIPAA.NPP, PRINT.DISABLED.DEP.LETTERS, COBRAN (indirectly)")

    t = doc.add_table(rows=7, cols=2)
    t.rows[0].cells[0].text = "Field Name"
    t.rows[0].cells[1].text = "Purpose"
    dp_fields = [
        ("DPDEP.NEED", "Comma-delimited codes: RR/RC/IA/LA/DL/NR"),
        ("DPDDL.SENT", "Multi-valued dates when disabled dep letters were sent"),
        ("DPDDL.LETTERS", "Multi-valued letter type codes sent on each date"),
        ("DPDISABILITY.EXPIRATION.DATE", "Required for Limited Approval (LA) letters"),
        ("DPREL.TYPE.EFFDATE", "Relationship type effective date"),
        ("DPADD1–DPZIP/DPCOUNTRY", "Dependent alternate address fields"),
    ]
    for i, (f, p) in enumerate(dp_fields, 1):
        t.rows[i].cells[0].text = f
        t.rows[i].cells[1].text = p
    style_table(t)

    add_code_block(doc,
        "CALL *READREC('',MAT DPREC,DEPENDENTS,EEID.DEPNO,'SKIP')\n"
        "CALL *WRITEREC('',MAT DPREC,DEPENDENTS,EEID.DEPNO,'TRANS.LOG,':PROGRAM.NAME)"
    )

    # 2.5 ERADJ
    doc.add_heading("2.5 ERADJ (Employer Adjustments) Record", level=2)
    doc.add_paragraph("Description: Billing adjustment records used to charge employers for COBRA and HIPAA notice fees.")
    doc.add_paragraph("Key: FILENO + BILLPERIOD + SUFFIX (e.g., 12345200401)")
    doc.add_paragraph("Loaded via: $INCLUDE BP ERADJ.INS → MAT EAREC")
    doc.add_paragraph("Used by: COBRAN (COBRACHG), HIPAA.NPP (HIPAA)")

    t = doc.add_table(rows=7, cols=2)
    t.rows[0].cells[0].text = "Field Name"
    t.rows[0].cells[1].text = "Purpose"
    ea_fields = [
        ("EADESC", "Description (e.g., '5 COBRA NOTICES')"),
        ("EACOMMENTS", "Multi-valued list of EEID/EEID.DEPNO processed"),
        ("EADISTFIELD", "Distribution field: COBRACHG or HIPAA"),
        ("EADISTAMOUNT", "Running total amount for this distribution line"),
        ("EACOVPERIOD", "Coverage period (YYYYMM)"),
        ("EAAMOUNT", "Grand total amount on the adjustment record"),
    ]
    for i, (f, p) in enumerate(ea_fields, 1):
        t.rows[i].cells[0].text = f
        t.rows[i].cells[1].text = p
    style_table(t)

    # 2.6 CORRESP.DATA.REC
    doc.add_heading("2.6 CORRESP.DATA.REC (In-Memory Dynamic Array)", level=2)
    doc.add_paragraph("Not a file — assembled in memory before calling *CREATE.CORRESP.TRACKING. Structure defined by CORRESP.INPUT.DEF record BASIC.SMARTCOMM.DOC.1.")
    doc.add_paragraph("Used by: All 3 programs")

    t = doc.add_table(rows=4, cols=3)
    t.rows[0].cells[0].text = "Position"
    t.rows[0].cells[1].text = "Name"
    t.rows[0].cells[2].text = "Purpose"
    cdr_fields = [
        ("<2>", "GroupNum / FILENO", "Employer group number"),
        ("<5>", "AdditionalEEIDDepno", "Additional dependent ID"),
        ("<6>", "AdditionalDependentPurpose", "RECIPIENT or SUBJECT"),
    ]
    for i, (a, b, c) in enumerate(cdr_fields, 1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
        t.rows[i].cells[2].text = c
    style_table(t)

    # 2.7 CORRESP.TRACKING
    doc.add_heading("2.7 CORRESP.TRACKING Record", level=2)
    doc.add_paragraph("Auto-generated by *CREATE.CORRESP.TRACKING. Interface to the SmartComm document generation system.")

    t = doc.add_table(rows=5, cols=2)
    t.rows[0].cells[0].text = "Parameter"
    t.rows[0].cells[1].text = "Value"
    ct_fields = [
        ("CORRESP.INPUT.DEF.ID", "'BASIC.SMARTCOMM.DOC.1'"),
        ("RECIPIENT.TYPE", "'DEPENDENT'"),
        ("METHOD.REQUESTED", "'LOOKUP'"),
        ("REQUEST.SOURCE", "PROGRAM.NAME"),
    ]
    for i, (a, b) in enumerate(ct_fields, 1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
    style_table(t)

    # 2.8 SAVEDLISTS
    doc.add_heading("2.8 SAVEDLISTS / Audit Lists", level=2)
    doc.add_paragraph("Named lists stored in the &SAVEDLISTS& system file for audit trails and re-run support.")

    t = doc.add_table(rows=9, cols=3)
    t.rows[0].cells[0].text = "Pattern"
    t.rows[0].cells[1].text = "Program"
    t.rows[0].cells[2].text = "Purpose"
    sl_data = [
        ("COBRAN.DONE{logname}", "COBRAN", "Primary completion list"),
        ("COBRAN.DONE{logname}.{date}.{time}", "COBRAN", "Timestamped completion list"),
        ("STORE.MERITAIN.COBRA.INTRO.LETTER", "COBRAN", "Employee notice tracking"),
        ("STORE.MERITAIN.COBRA.DEPENDENT.NOTICE", "COBRAN", "Dependent notice tracking"),
        ("COBRA.DEP.NOTICE.{logname}", "COBRAN", "Dependent EEID.DEPNO list"),
        ("HIPAA.PDONE{logname}", "HIPAA.NPP", "All processed EEID.DEPNOs"),
        ("HIPAA.PDONE.SC.{logname}", "HIPAA.NPP", "SmartComm-routed (SF) members"),
        ("EEID.LIST.{logname}", "HIPAA.NPP", "Filtered EEID list"),
    ]
    for i, (a, b, c) in enumerate(sl_data, 1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
        t.rows[i].cells[2].text = c
    style_table(t)

    # 2.9 Supplementary Records
    doc.add_heading("2.9 Supplementary Records", level=2)
    doc.add_paragraph("These records are used primarily by HIPAA.NPP for the Fully-Insured (FI) PCL generation path:")

    t = doc.add_table(rows=9, cols=2)
    t.rows[0].cells[0].text = "Record"
    t.rows[0].cells[1].text = "Purpose"
    supp_data = [
        ("SYSMGMT (HIPAA.DAILY.COUNTER)", "Daily run counter"),
        ("GROUPINFO", "Group-level NPP configuration (GINPP flag)"),
        ("CERTS.LAYOUT", "FI form layout lookup by insurance carrier"),
        ("CERTS.MERGE", "Merge field definitions for FI forms"),
        ("CERTS.TEXT2", "Text content for FI HIPAA form pages"),
        ("LETTER.COPIES", "Copy count configuration for FI printed notices"),
        ("OUTPUT.CONTROL", "FI output config including HIPAA.YEARS"),
        ("FTP.LOG / FTP.REQUESTS", "FTP transmission tracking for RedCard print vendor"),
    ]
    for i, (a, b) in enumerate(supp_data, 1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
    style_table(t)

    # 2.10 SmartComm Configuration
    doc.add_heading("2.10 SmartComm Configuration", level=2)
    doc.add_paragraph("All 9 document types use the BASIC.SMARTCOMM.DOC.1 setup in CORRESP.INPUT.DEF:")

    t = doc.add_table(rows=10, cols=6)
    headers = ["#", "Letter Description", "Document Name", "SmartComm ID", "Returns Image", "Program"]
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    sc_data = [
        ("1", "HIPAA Notice of Privacy Practices", "Meritain.HIPAA.Notice", "690679345", "N", "HIPAA.NPP"),
        ("2", "COBRA Introduction Letter", "Meritain.COBRA.Intro.Letter", "690680007", "N", "COBRAN"),
        ("3", "Dependent COBRA Introduction Letter", "Meritain.COBRA.Dependent.Notice", "690680105", "N", "COBRAN"),
        ("4", "Disabled Dep Denial Notice", "Disabled.Dependent.-.Denial", "690918681", "Y", "PDDL"),
        ("5", "Disabled Dep Indefinite Approval", "Disabled.Dependent.-.Indefinite.Approval", "690918738", "Y", "PDDL"),
        ("6", "Disabled Dep Recertification", "Disabled.Dependent.-.Recertification", "690918752", "Y", "PDDL"),
        ("7", "Disabled Dep Review Request", "Disabled.Dependent.-.Review.Request", "690918765", "Y", "PDDL"),
        ("8", "Disabled Dep Limited Approval", "Disabled.Dependent.-.Limited.Approval", "690918728", "Y", "PDDL"),
        ("9", "Disabled Dep No Response Letter", "Disabled.Dependent.-.No.Response", "690918705", "Y", "PDDL"),
    ]
    for i, row_data in enumerate(sc_data, 1):
        for j, val in enumerate(row_data):
            t.rows[i].cells[j].text = val
    style_table(t)

    # ======================================================================
    # SECTION 3: THE 9 LETTER TYPES
    # ======================================================================
    doc.add_heading("3. The 9 Letter Types — Identified from Code", level=1)

    doc.add_paragraph("Each letter type is discoverable directly from the source code:")

    doc.add_heading("From COBRAN.md", level=3)
    doc.add_paragraph("1. Meritain.COBRA.Intro.Letter — Assigned to TEMPLATE.NAME at line 178", style="List Bullet")
    doc.add_paragraph("2. Meritain.COBRA.Dependent.Notice — Assigned to DEP.TEMPLATE.NAME at line 182", style="List Bullet")

    doc.add_heading("From HIPAA.NPP.md", level=3)
    doc.add_paragraph("3. Meritain.HIPAA.Notice — Assigned to DOCUMENT.NAME at line 1287", style="List Bullet")

    doc.add_heading("From PRINT.DISABLED.DEP.LETTERS.md (lines 148–173)", level=3)
    doc.add_paragraph("4. Disabled.Dependent.-.Review.Request (code: RR)", style="List Bullet")
    doc.add_paragraph("5. Disabled.Dependent.-.Recertification (code: RC)", style="List Bullet")
    doc.add_paragraph("6. Disabled.Dependent.-.Indefinite.Approval (code: IA)", style="List Bullet")
    doc.add_paragraph("7. Disabled.Dependent.-.Limited.Approval (code: LA)", style="List Bullet")
    doc.add_paragraph("8. Disabled.Dependent.-.Denial (code: DL)", style="List Bullet")
    doc.add_paragraph("9. Disabled.Dependent.-.No.Response (code: NR)", style="List Bullet")

    doc.add_heading("Summary Table", level=3)
    t = doc.add_table(rows=10, cols=7)
    headers = ["#", "Letter Name", "Template/Document Name", "Source Program", "Trigger Condition", "SmartComm ID", "Returns Image"]
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    letter_data = [
        ("1", "COBRA Intro Letter", "Meritain.COBRA.Intro.Letter", "COBRAN", "EENEED='C' + ERCOBRAFEE>0", "690680007", "N"),
        ("2", "COBRA Dep Notice", "Meritain.COBRA.Dependent.Notice", "COBRAN", "Spouse at diff addr", "690680105", "N"),
        ("3", "HIPAA Privacy Notice", "Meritain.HIPAA.Notice", "HIPAA.NPP", "EENEED='N' + SF group", "690679345", "N"),
        ("4", "DD Review Request", "Disabled.Dependent.-.Review.Request", "PDDL", "DPDEP.NEED has 'RR'", "690918765", "Y"),
        ("5", "DD Recertification", "Disabled.Dependent.-.Recertification", "PDDL", "DPDEP.NEED has 'RC'", "690918752", "Y"),
        ("6", "DD Indef. Approval", "Disabled.Dependent.-.Indefinite.Approval", "PDDL", "DPDEP.NEED has 'IA'", "690918738", "Y"),
        ("7", "DD Limited Approval", "Disabled.Dependent.-.Limited.Approval", "PDDL", "DPDEP.NEED has 'LA'", "690918728", "Y"),
        ("8", "DD Denial", "Disabled.Dependent.-.Denial", "PDDL", "DPDEP.NEED has 'DL'", "690918681", "Y"),
        ("9", "DD No Response", "Disabled.Dependent.-.No.Response", "PDDL", "DPDEP.NEED has 'NR'", "690918705", "Y"),
    ]
    for i, row_data in enumerate(letter_data, 1):
        for j, val in enumerate(row_data):
            t.rows[i].cells[j].text = val
    style_table(t)

    # ======================================================================
    # SECTION 4: DEPENDENCY GRAPH
    # ======================================================================
    doc.add_heading("4. Dependency Graph", level=1)
    doc.add_paragraph("The following diagram shows how the 10+ record types relate to each other:")

    mermaid_dep = """graph TD
    subgraph "Primary Data Records"
        EE["EMPLOYEES"]
        ER["EMPLOYERS"]
        EO["EROPTIONS"]
        DEP["DEPENDENTS"]
    end
    subgraph "Billing Records"
        ERADJ["ERADJ"]
    end
    subgraph "Correspondence Records"
        CDR["CORRESP.DATA.REC"]
        CT["CORRESP.TRACKING"]
    end
    subgraph "Audit Records"
        SL["SAVEDLISTS"]
        SM["SYSMGMT"]
    end
    EE -- "EEFILENO" --> ER
    EE -- "EEFILENO" --> EO
    EE -- "EEID.DEPNO" --> DEP
    ER -- "ERBILLEDTO" --> ERADJ
    EE --> CDR
    ER --> CDR
    DEP --> CDR
    CDR --> CT"""
    add_mermaid_diagram(doc, mermaid_dep, "Figure 1: Record Type Dependency Graph")

    # ======================================================================
    # SECTION 5: COUPLING MAP
    # ======================================================================
    doc.add_heading("5. Current Coupling Map", level=1)
    doc.add_paragraph("This diagram shows the 8+ functional concerns tangled across the 3 programs:")

    mermaid_coupling = """graph LR
    subgraph "COBRAN"
        C1["Eligibility Filtering"]
        C2["Employer Loading"]
        C3["Letter Gen (EE)"]
        C4["Letter Gen (Dep)"]
        C5["Billing"]
        C6["Audit Lists"]
    end
    subgraph "HIPAA.NPP"
        H1["Eligibility Filtering"]
        H2["Employer Loading"]
        H3["SF/FI Routing"]
        H4["SmartComm"]
        H5["Legacy PCL"]
        H6["Billing"]
        H7["FTP"]
        H8["Audit Lists"]
    end
    subgraph "PDDL"
        P1["Eligibility Filtering"]
        P2["Dep Validation"]
        P3["Letter Gen"]
        P4["Dep Update"]
        P5["Error Report"]
    end"""
    add_mermaid_diagram(doc, mermaid_coupling, "Figure 2: Current Coupling Map")

    # ======================================================================
    # SECTION 6: COUPLING POINTS TABLE
    # ======================================================================
    doc.add_heading("6. Coupling Points Table", level=1)

    t = doc.add_table(rows=9, cols=3)
    t.rows[0].cells[0].text = "Concern"
    t.rows[0].cells[1].text = "Where It's Tangled"
    t.rows[0].cells[2].text = "Shared Records"
    coupling_data = [
        ("1. Eligibility Filtering", "COBRAN MAIN.LOOP, HIPAA.NPP FILTER.EEIDS, PDDL VALIDATE.DEP.DATA", "EMPLOYEES"),
        ("2. Employer Context Loading", "COBRAN MAIN.LOOP (lines 232-260), HIPAA.NPP SET.SKIP.FILENO", "EMPLOYERS, EROPTIONS"),
        ("3. SmartComm Submission", "All 3: SUBMIT.CORRESP.REQUEST", "CORRESP.DATA.REC, CORRESP.TRACKING"),
        ("4. Billing (ERADJ)", "COBRAN DO.BILLING (COBRACHG), HIPAA.NPP DO.BILLING (HIPAA)", "ERADJ"),
        ("5. FI HIPAA Generation", "HIPAA.NPP only: GET/GENERATE.HIPAA.FORMS", "CERTS.LAYOUT/MERGE/TEXT2"),
        ("6. FTP Transmission", "HIPAA.NPP only: FTP.FILE", "FTP.LOG, FTP.REQUESTS"),
        ("7. Audit List Management", "COBRAN END.REPORT.FILE, HIPAA.NPP inline SAVE.LIST", "&SAVEDLISTS&"),
        ("8. Dep Record Update", "PDDL: WRITE.DEP.REC + UPDATE.DEP.NEED", "DEPENDENTS"),
    ]
    for i, (a, b, c) in enumerate(coupling_data, 1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
        t.rows[i].cells[2].text = c
    style_table(t)

    # ======================================================================
    # SECTION 7: DECOUPLING STRATEGY
    # ======================================================================
    doc.add_heading("7. Decoupling Strategy", level=1)
    doc.add_paragraph("The following 8 extractions would convert duplicated logic into shared, cataloged subroutines:")

    extractions = [
        ("7.1 *CHECK.LETTER.ELIGIBILITY",
         "SUBROUTINE *CHECK.LETTER.ELIGIBILITY(IS.ELIGIBLE, MAT EEREC, EEID, TODAY, LETTER.TYPE, OPTIONS)",
         "COBRAN MAIN.LOOP, HIPAA.NPP FILTER.EEIDS, PDDL VALIDATE.DEP.DATA",
         "Consolidates EENEED check, effective date location, and status filtering into a single subroutine."),
        ("7.2 *LOAD.EMPLOYER.CONTEXT",
         "SUBROUTINE *LOAD.EMPLOYER.CONTEXT(LOAD.STATUS, MAT ERREC, MAT EOREC, FILENO, LAST.FILENO, TODAY)",
         "COBRAN MAIN.LOOP (lines 232-260), HIPAA.NPP SET.SKIP.FILENO",
         "Wraps the cached employer+eroptions read into a single subroutine."),
        ("7.3 *SUBMIT.LETTER.REQUEST",
         "SUBROUTINE *SUBMIT.LETTER.REQUEST(ERROR.MESSAGES, DOCUMENT.NAME, RECIPIENT.ID, FILENO, ...)",
         "All 3 programs: SUBMIT.CORRESP.REQUEST",
         "Wraps CORRESP.DATA.REC assembly and *CREATE.CORRESP.TRACKING call."),
        ("7.4 *ADD.NOTICE.BILLING",
         "SUBROUTINE *ADD.NOTICE.BILLING(MAT EAREC, FILENO, EEID, ..., DIST.FIELD, FEE.AMOUNT)",
         "COBRAN DO.BILLING (COBRACHG), HIPAA.NPP DO.BILLING (HIPAA)",
         "Parameterizes the billing subroutine with DIST.FIELD and FEE.AMOUNT."),
        ("7.5 *GENERATE.FI.HIPAA.NOTICE",
         "SUBROUTINE *GENERATE.FI.HIPAA.NOTICE(OUTPUT.STATUS, EEID, FILENO, ...)",
         "HIPAA.NPP FI branch: GET.HIPAA.FORMS + GENERATE.HIPAA.FORMS",
         "Encapsulates the entire FI PCL generation path."),
        ("7.6 *SUBMIT.PRINT.FTP",
         "SUBROUTINE *SUBMIT.PRINT.FTP(FTP.STATUS, FILENAME, FTP.ID, PROGRAM.NAME)",
         "HIPAA.NPP FTP.FILE",
         "Wraps FTP submission logic for reuse."),
        ("7.7 *SAVE.PROCESS.AUDIT",
         "SUBROUTINE *SAVE.PROCESS.AUDIT(LIST.DATA, LIST.NAME.PREFIX, UV.LOGNAME, INCLUDE.TIMESTAMP)",
         "COBRAN END.REPORT.FILE, HIPAA.NPP inline SAVE.LIST commands",
         "Generalizes audit list saving pattern."),
        ("7.8 *UPDATE.DEP.LETTER.STATUS",
         "SUBROUTINE *UPDATE.DEP.LETTER.STATUS(MAT DPREC, EEID.DEPNO, CUR.NEED, TODAY, ...)",
         "PDDL WRITE.DEP.REC + UPDATE.DEP.NEED",
         "Combines letter-sent tracking and dependent record write-back."),
    ]
    for title, sig, source, desc in extractions:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Proposed Signature:")
        add_code_block(doc, sig)
        doc.add_paragraph(f"Extracts from: {source}")
        doc.add_paragraph(desc)

    # ======================================================================
    # SECTION 8: DECOUPLED ARCHITECTURE
    # ======================================================================
    doc.add_heading("8. Decoupled Architecture", level=1)
    doc.add_paragraph("After extracting the 8 shared subroutines, each program becomes a thin orchestrator:")

    mermaid_decoupled = """graph TD
    COBRAN["COBRAN (orchestrator)"]
    HIPAA["HIPAA.NPP (orchestrator)"]
    PDDL["PDDL (orchestrator)"]
    EF["*CHECK.LETTER.ELIGIBILITY"]
    EC["*LOAD.EMPLOYER.CONTEXT"]
    CS["*SUBMIT.LETTER.REQUEST"]
    BA["*ADD.NOTICE.BILLING"]
    FI["*GENERATE.FI.HIPAA.NOTICE"]
    FTP["*SUBMIT.PRINT.FTP"]
    AL["*SAVE.PROCESS.AUDIT"]
    DU["*UPDATE.DEP.LETTER.STATUS"]
    COBRAN --> EF
    COBRAN --> EC
    COBRAN --> CS
    COBRAN --> BA
    COBRAN --> AL
    HIPAA --> EF
    HIPAA --> EC
    HIPAA --> CS
    HIPAA --> BA
    HIPAA --> FI
    HIPAA --> FTP
    HIPAA --> AL
    PDDL --> EF
    PDDL --> CS
    PDDL --> DU
    PDDL --> EC"""
    add_mermaid_diagram(doc, mermaid_decoupled, "Figure 3: Decoupled Architecture")

    # ======================================================================
    # SECTION 9: PRIORITY ORDER
    # ======================================================================
    doc.add_heading("9. Priority Order for Decoupling", level=1)

    t = doc.add_table(rows=9, cols=4)
    t.rows[0].cells[0].text = "Priority"
    t.rows[0].cells[1].text = "Extraction"
    t.rows[0].cells[2].text = "Impact"
    t.rows[0].cells[3].text = "Reason"
    priority_data = [
        ("1", "*SUBMIT.LETTER.REQUEST", "All 3 programs", "Highest duplication — identical pattern in all 3 programs"),
        ("2", "*CHECK.LETTER.ELIGIBILITY", "All 3 programs", "Core eligibility logic duplicated with subtle variations"),
        ("3", "*LOAD.EMPLOYER.CONTEXT", "COBRAN + HIPAA.NPP", "Identical cached read pattern must stay in sync"),
        ("4", "*ADD.NOTICE.BILLING", "COBRAN + HIPAA.NPP", "ERADJ loop identical except EADISTFIELD value"),
        ("5", "*SAVE.PROCESS.AUDIT", "COBRAN + HIPAA.NPP", "Audit list conventions should be consistent"),
        ("6", "*UPDATE.DEP.LETTER.STATUS", "PDDL only", "Complex dep record update; improves testability"),
        ("7", "*GENERATE.FI.HIPAA.NOTICE", "HIPAA.NPP only", "Isolates legacy PCL path for future retirement"),
        ("8", "*SUBMIT.PRINT.FTP", "HIPAA.NPP only", "FTP logic self-contained; enables future reuse"),
    ]
    for i, (a, b, c, d) in enumerate(priority_data, 1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
        t.rows[i].cells[2].text = c
        t.rows[i].cells[3].text = d
    style_table(t)

    # ======================================================================
    # SECTION 10: PROGRAM FLOW DIAGRAMS
    # ======================================================================
    doc.add_heading("10. Program Flow Diagrams", level=1)

    # COBRAN Flow
    doc.add_heading("COBRAN Flow", level=2)
    mermaid_cobran = """graph TD
    START["COBRAN Start"] --> INIT["INIT.VARS"]
    INIT --> SELECT["SELECT EMPLOYEES NEED='C'"]
    SELECT --> LOOP["MAIN.LOOP: READNEXT EEID"]
    LOOP --> READ_EE["Read EMPLOYEES record"]
    READ_EE --> CHECK["Check EENEED for 'X'"]
    CHECK --> LOCATE["LOCATE effective date index"]
    LOCATE --> LOAD_ER["Load EMPLOYERS + EROPTIONS"]
    LOAD_ER --> STATUS["Check status, COBRAFEE"]
    STATUS --> DEP["Check dep COBRA eligibility"]
    DEP -->|"Spouse"| SUBMIT_DP["SUBMIT.CORRESP.REQUEST (DP)"]
    DEP -->|"Employee"| SUBMIT_EE["SUBMIT.CORRESP.REQUEST (EE)"]
    SUBMIT_DP --> BILLING["DO.BILLING (COBRACHG)"]
    SUBMIT_EE --> BILLING
    BILLING --> LOOP
    LOOP -->|"End"| AUDIT["END.REPORT.FILE"]"""
    add_mermaid_diagram(doc, mermaid_cobran, "Figure 4: COBRAN Program Flow")

    # HIPAA.NPP Flow
    doc.add_heading("HIPAA.NPP Flow", level=2)
    mermaid_hipaa = """graph TD
    START["HIPAA.NPP Start"] --> INIT["Initialize + COUNTER"]
    INIT --> FILTER["FILTER.EEIDS"]
    FILTER --> SF_FI{"SF or FI?"}
    SF_FI -->|"Self-Funded"| SF["SF Loop"]
    SF_FI -->|"Fully-Insured"| FI["FI Loop"]
    SF --> SC["SUBMIT.CORRESP.REQUEST"]
    SC --> SFB["DO.BILLING (HIPAA)"]
    SFB --> SF
    FI --> GET["GET.HIPAA.FORMS"]
    GET --> GEN["GENERATE.HIPAA.FORMS"]
    GEN --> FIB["DO.BILLING (HIPAA)"]
    FIB --> FI
    FI -->|"Done"| FTP["FTP.FILE"]
    SF -->|"Done"| SAVE["Save audit lists"]
    FTP --> SAVE"""
    add_mermaid_diagram(doc, mermaid_hipaa, "Figure 5: HIPAA.NPP Program Flow")

    # PDDL Flow
    doc.add_heading("PRINT.DISABLED.DEP.LETTERS Flow", level=2)
    mermaid_pddl = """graph TD
    START["PDDL Start"] --> INIT["INIT.VARS (6 letter types)"]
    INIT --> SELECT["SELECT EMPLOYEES STATUS='H'"]
    SELECT --> LOOP["MAIN.LOOP: READNEXT EEID"]
    LOOP --> VALIDATE["VALIDATE.DEP.DATA"]
    VALIDATE --> CHECK["CHECK.DEP.NEED"]
    CHECK --> NEED["For each CUR.NEED"]
    NEED --> SET["SET.SMARTCOMM.TEMPLATE"]
    SET --> SUBMIT["SUBMIT.CORRESP.REQUEST"]
    SUBMIT -->|"Success"| UPDATE["UPDATE.DEP.NEED"]
    SUBMIT -->|"Error"| ERR["ERROR.DEP.NEED"]
    UPDATE --> NEED
    ERR --> NEED
    NEED -->|"Done"| WRITE["WRITE.DEP.REC"]
    WRITE --> LOOP
    LOOP -->|"End"| EMAIL["Email error report"]"""
    add_mermaid_diagram(doc, mermaid_pddl, "Figure 6: PRINT.DISABLED.DEP.LETTERS Program Flow")

    # ======================================================================
    # SAVE
    # ======================================================================
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "DG_Analysis_Report.docx")
    doc.save(output_path)
    print(f"Report saved to {output_path}")
    return output_path


if __name__ == "__main__":
    build_report()
